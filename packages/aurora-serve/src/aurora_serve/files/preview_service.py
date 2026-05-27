"""Office file preview service - converts Office files to previewable formats."""

from __future__ import annotations

import hashlib
import logging
import shutil
import subprocess
import tempfile
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Literal

from aurora_serve.excel.reader import ExcelReader
from aurora_serve.metadata import storage_dir

logger = logging.getLogger(__name__)

# LibreOffice soffice path - check common locations
SOFFICE_PATHS = [
    "/opt/homebrew/bin/soffice",  # macOS Homebrew
    "/usr/bin/soffice",  # Linux
    "/Applications/LibreOffice.app/Contents/MacOS/soffice",  # macOS App
]

PreviewFormat = Literal["html", "pdf", "metadata"]


@dataclass
class PreviewResult:
    """Result of a preview conversion."""
    kind: str  # "spreadsheet", "document", "presentation"
    format: PreviewFormat
    content: str | bytes  # HTML string or PDF bytes
    metadata: dict[str, Any] | None = None


def find_soffice() -> str | None:
    """Find LibreOffice soffice binary."""
    for path in SOFFICE_PATHS:
        if Path(path).exists():
            return path
    # Also check PATH
    return shutil.which("soffice")


def get_cache_dir() -> Path:
    """Get or create the preview cache directory."""
    cache_dir = storage_dir() / "preview_cache"
    cache_dir.mkdir(parents=True, exist_ok=True)
    return cache_dir


def get_cache_key(file_path: Path) -> str:
    """Generate a cache key from file path and modification time."""
    mtime = file_path.stat().st_mtime
    key = f"{file_path.name}_{mtime}"
    return hashlib.md5(key.encode()).hexdigest()


class OfficePreviewService:
    """Handles Office file preview conversion."""

    def __init__(self):
        self._soffice_path = find_soffice()
        self._cache_dir = get_cache_dir()
        self._soffice_available: bool | None = None  # cached result

    def is_soffice_available(self) -> bool:
        """Check if LibreOffice is available for PDF conversion.

        Result is cached after the first call to avoid blocking every request
        with a subprocess invocation (which can take up to 10 seconds on timeout).
        """
        if self._soffice_available is not None:
            return self._soffice_available
        if self._soffice_path is None:
            self._soffice_available = False
            return False
        try:
            result = subprocess.run(
                [self._soffice_path, "--version"],
                capture_output=True,
                timeout=10,
            )
            self._soffice_available = result.returncode == 0
        except (subprocess.TimeoutExpired, FileNotFoundError, OSError):
            self._soffice_available = False
        return self._soffice_available

    def preview_excel(self, file_path: Path) -> PreviewResult:
        """Generate HTML preview for Excel files using ExcelReader."""
        try:
            with ExcelReader(str(file_path)) as reader:
                # Get column names and metadata
                col_names, col_data = reader.get_columns()
                columns = [
                    {
                        "name": row[0],
                        "type": row[1],
                        "nullable": row[2] == "YES",
                    }
                    for row in col_data
                ]

                # Get sample data (first 100 rows)
                temp_table = reader.temp_table
                sample_sql = f"SELECT * FROM {temp_table} LIMIT 100"
                _, sample_rows = reader.run_sql(sample_sql)

                # Get row count
                count_sql = f"SELECT COUNT(*) FROM {temp_table}"
                _, count_result = reader.run_sql(count_sql)
                total_rows = count_result[0][0] if count_result else 0

                # Build HTML
                html = self._build_excel_html(col_names, sample_rows, columns, total_rows)

                return PreviewResult(
                    kind="spreadsheet",
                    format="html",
                    content=html,
                    metadata={
                        "columns": columns,
                        "sampleRows": len(sample_rows),
                        "totalRows": total_rows,
                    },
                )
        except Exception as e:
            logger.error(f"Excel preview error: {e}")
            return PreviewResult(
                kind="spreadsheet",
                format="metadata",
                content="<p>Preview unavailable. Download file to view.</p>",
                metadata={"error": str(e)},
            )

    def _build_excel_html(
        self,
        columns: list[str],
        rows: list[tuple[Any, ...]],
        column_meta: list[dict],
        total_rows: int,
    ) -> str:
        """Build HTML table from Excel data with Tailwind styling."""
        # Header
        html_parts = [
            '<div class="p-4">',
            f'<p class="text-sm text-gray-500 mb-2">Showing {len(rows)} of {total_rows} rows</p>',
            '<div class="overflow-auto max-h-[600px] border rounded-lg">',
            '<table class="min-w-full divide-y divide-gray-200">',
        ]

        # Table header
        html_parts.append('<thead class="bg-gray-50 sticky top-0">')
        html_parts.append("<tr>")
        for i, col in enumerate(columns):
            col_type = column_meta[i]["type"] if i < len(column_meta) else "unknown"
            html_parts.append(
                f'<th class="px-4 py-3 text-left text-xs font-medium text-gray-500 uppercase tracking-wider whitespace-nowrap">'
                f'<div>{col}</div>'
                f'<div class="text-xs text-gray-400 normal-case">{col_type}</div>'
                f"</th>"
            )
        html_parts.append("</tr>")
        html_parts.append("</thead>")

        # Table body
        html_parts.append('<tbody class="bg-white divide-y divide-gray-200">')
        for row in rows:
            html_parts.append("<tr>")
            for cell in row:
                # Format cell value
                cell_str = str(cell) if cell is not None else ""
                # Escape HTML
                cell_escaped = cell_str.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                html_parts.append(
                    f'<td class="px-4 py-2 text-sm text-gray-900 whitespace-nowrap">{cell_escaped}</td>'
                )
            html_parts.append("</tr>")
        html_parts.append("</tbody>")

        # Footer
        html_parts.extend([
            "</table>",
            "</div>",
            "</div>",
        ])

        return "\n".join(html_parts)

    def preview_document(self, file_path: Path) -> PreviewResult:
        """Convert DOCX to PDF using soffice."""
        if not self.is_soffice_available():
            return self._fallback_preview(file_path, "document")

        pdf_bytes = self._convert_to_pdf(file_path)
        if pdf_bytes is None:
            return self._fallback_preview(file_path, "document")

        metadata = self._extract_docx_metadata(file_path)
        return PreviewResult(
            kind="document",
            format="pdf",
            content=pdf_bytes,
            metadata=metadata,
        )

    def preview_presentation(self, file_path: Path) -> PreviewResult:
        """Convert PPTX to PDF using soffice."""
        if not self.is_soffice_available():
            return self._fallback_preview(file_path, "presentation")

        pdf_bytes = self._convert_to_pdf(file_path)
        if pdf_bytes is None:
            return self._fallback_preview(file_path, "presentation")

        metadata = self._extract_pptx_metadata(file_path)
        return PreviewResult(
            kind="presentation",
            format="pdf",
            content=pdf_bytes,
            metadata=metadata,
        )

    def _convert_to_pdf(self, file_path: Path) -> bytes | None:
        """Use soffice to convert Office file to PDF with caching."""
        # Check cache first
        cache_key = get_cache_key(file_path)
        cache_path = self._cache_dir / f"{cache_key}.pdf"

        if cache_path.exists():
            logger.info(f"Using cached PDF for {file_path.name}")
            return cache_path.read_bytes()

        # Create temp output directory
        with tempfile.TemporaryDirectory() as tmpdir:
            cmd = [
                self._soffice_path,
                "--headless",
                "--convert-to", "pdf",
                "--outdir", tmpdir,
                str(file_path),
            ]

            try:
                result = subprocess.run(
                    cmd,
                    check=True,
                    timeout=60,
                    capture_output=True,
                    text=True,
                )
                logger.debug(f"soffice output: {result.stdout}")

                # Find generated PDF
                pdf_name = file_path.stem + ".pdf"
                pdf_path = Path(tmpdir) / pdf_name

                if not pdf_path.exists():
                    logger.error(f"PDF not found at {pdf_path}")
                    return None

                pdf_bytes = pdf_path.read_bytes()

                # Cache the result
                cache_path.write_bytes(pdf_bytes)
                logger.info(f"Cached PDF for {file_path.name}")

                return pdf_bytes

            except subprocess.TimeoutExpired:
                logger.error(f"soffice timeout for {file_path.name}")
                return None
            except subprocess.CalledProcessError as e:
                logger.error(f"soffice error: {e.stderr}")
                return None
            except Exception as e:
                logger.error(f"PDF conversion error: {e}")
                return None

    def _extract_docx_metadata(self, file_path: Path) -> dict[str, Any]:
        """Extract metadata from DOCX using python-docx."""
        try:
            from docx import Document
            doc = Document(str(file_path))
            return {
                "paragraphs": len(doc.paragraphs),
                "tables": len(doc.tables),
                "sections": len(doc.sections),
            }
        except ImportError:
            logger.debug("python-docx not installed, skipping metadata extraction")
            return {}
        except Exception as e:
            logger.debug(f"DOCX metadata extraction error: {e}")
            return {}

    def _extract_pptx_metadata(self, file_path: Path) -> dict[str, Any]:
        """Extract metadata from PPTX using python-pptx."""
        try:
            from pptx import Presentation
            prs = Presentation(str(file_path))
            return {
                "slides": len(prs.slides),
                "slideWidth": prs.slide_width,
                "slideHeight": prs.slide_height,
            }
        except ImportError:
            logger.debug("python-pptx not installed, skipping metadata extraction")
            return {}
        except Exception as e:
            logger.debug(f"PPTX metadata extraction error: {e}")
            return {}

    def _fallback_preview(self, file_path: Path, kind: str) -> PreviewResult:
        """Fallback preview when soffice is unavailable."""
        metadata = {}
        if kind == "document":
            metadata = self._extract_docx_metadata(file_path)
        elif kind == "presentation":
            metadata = self._extract_pptx_metadata(file_path)

        return PreviewResult(
            kind=kind,
            format="metadata",
            content="",
            metadata={
                **metadata,
                "previewAvailable": False,
                "message": "Preview requires LibreOffice. Download file to view.",
            },
        )

    def clear_cache(self) -> None:
        """Clear all cached preview files."""
        if self._cache_dir.exists():
            for file in self._cache_dir.iterdir():
                file.unlink()
            logger.info("Preview cache cleared")


# Singleton instance
_preview_service: OfficePreviewService | None = None


def get_preview_service() -> OfficePreviewService:
    """Get the singleton preview service instance."""
    global _preview_service
    if _preview_service is None:
        _preview_service = OfficePreviewService()
    return _preview_service