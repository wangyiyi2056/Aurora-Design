"""DOCX parser using python-docx.

Migrated from LightRAG native DOCX parser.  Extracts paragraphs and
tables in document order, using tab-delimited format for table cells.
"""

from __future__ import annotations

import logging
from pathlib import Path
from typing import Any

from aurora_ext.rag.parser.base import BaseParser, ParseResult

logger = logging.getLogger(__name__)


class DocxParser(BaseParser):
    """Parser for .docx files using ``python-docx``."""

    _EXTENSIONS = {"docx", "odt"}

    @property
    def supported_extensions(self) -> set[str]:
        return self._EXTENSIONS

    async def parse(self, file_path: str | Path, **kwargs: Any) -> ParseResult:
        from docx import Document

        path = Path(file_path)

        try:
            doc = Document(str(path))
        except Exception as exc:
            raise ValueError(f"Cannot open DOCX: {path}") from exc

        parts: list[str] = []
        table_count = 0

        for element in doc.element.body:
            tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag

            if tag == "p":
                for para in doc.paragraphs:
                    if para._element is element:
                        text = para.text.strip()
                        if text:
                            parts.append(text)
                        break

            elif tag == "tbl":
                for table in doc.tables:
                    if table._element is element:
                        table_count += 1
                        table_parts: list[str] = []
                        for row in table.rows:
                            cells = [cell.text.strip() for cell in row.cells]
                            table_parts.append("\t".join(cells))
                        if table_parts:
                            parts.append("\n".join(table_parts))
                        break

        full_text = "\n\n".join(parts)

        return ParseResult(
            text=full_text,
            file_path=str(path),
            file_type=path.suffix.lstrip(".").lower(),
            metadata={
                "paragraph_count": len(doc.paragraphs),
                "table_count": table_count,
                "char_count": len(full_text),
            },
        )
